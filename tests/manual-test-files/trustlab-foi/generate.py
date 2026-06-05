#!/usr/bin/env python3
"""Generate the TrustLab Inc / CCSA FOI demo content.

Outputs (under tests/manual-test-files/trustlab-foi/generated/):

  trustlab-cv.docx                  CV for TrustLab Inc's principal
  trustlab-proposal-letter.docx     Proposal pitch (cover letter equivalent)
  ccsa-evaluation-memo.docx         Synthetic internal CCSA evaluation memo
  ccsa-referee-notes.docx           Synthetic internal CCSA reference-check notes
  email-thread/
    01-foi-request.eml              External requester -> CCSA FOI office
    02-acknowledgment.eml           CCSA FOI office -> requester
    03-internal-forward.eml         FOI officer -> Procurement records
    04-internal-response.eml        Procurement -> FOI officer (with attachments)

The DOCX files are rendered from the four markdown files in ./source/.

Run from the repo root:

    python3 tests/manual-test-files/trustlab-foi/generate.py

Or inside the backend container (python-docx is already installed there):

    docker compose exec backend python tests/manual-test-files/trustlab-foi/generate.py

The script is deterministic. Re-running produces byte-identical output for
the EMLs (same Message-IDs, same Date headers). DOCX bytes vary because
python-docx embeds a creation timestamp; that's fine for fixtures.

Privacy substitutions baked in:
  - Principal's real phone -> 555-0142 (fictional range, per existing
    fixture convention).
  - All third-party names (referees, committee members, FOI requester,
    FOI officer, records custodian) are fictional.
"""
from __future__ import annotations

import re
import sys
from datetime import datetime, timezone
from email.message import EmailMessage
from email.utils import formatdate
from pathlib import Path

# --- Paths -------------------------------------------------------------------

HERE = Path(__file__).parent
SOURCE = HERE / "source"
OUT = HERE / "generated"
THREAD_OUT = OUT / "email-thread"
OUT.mkdir(exist_ok=True)
THREAD_OUT.mkdir(exist_ok=True)


# --- Synthetic identities ----------------------------------------------------

REQUESTER = ("Jordan Park", "jordan.park@example.org")
FOI_OFFICER = ("Riley Chen", "riley.chen@ccsa-bc.example.org")
PROCUREMENT_RECORDS = ("Sam Beaumont", "sam.beaumont@ccsa-bc.example.org")


# --- Minimal markdown -> DOCX ------------------------------------------------

INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
INLINE_CODE = re.compile(r"`([^`]+)`")


def _add_runs(paragraph, text: str) -> None:
    """Walk inline bold/italic/code markers and add styled runs."""
    # Split on bold first, then italic within non-bold segments.
    pos = 0
    for m in INLINE_BOLD.finditer(text):
        if m.start() > pos:
            _add_italic_runs(paragraph, text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        _add_italic_runs(paragraph, text[pos:])


def _add_italic_runs(paragraph, text: str) -> None:
    pos = 0
    for m in INLINE_ITALIC.finditer(text):
        if m.start() > pos:
            _add_code_runs(paragraph, text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.italic = True
        pos = m.end()
    if pos < len(text):
        _add_code_runs(paragraph, text[pos:])


def _add_code_runs(paragraph, text: str) -> None:
    pos = 0
    for m in INLINE_CODE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.font.name = "Courier New"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    """Render a markdown file to a DOCX. Subset of markdown only — enough
    for the four seed documents in ./source/."""
    from docx import Document

    doc = Document()
    lines = md_path.read_text().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Blank line — paragraph break is implicit
        if not line:
            i += 1
            continue

        # Horizontal rule — render as a blank paragraph
        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
            i += 1
            continue

        # Block quote
        if line.startswith("> "):
            block = []
            while i < len(lines) and lines[i].rstrip().startswith(">"):
                block.append(lines[i].rstrip().lstrip("> ").lstrip(">"))
                i += 1
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs(p, " ".join(b.strip() for b in block))
            continue

        # Tables: a line that starts with "|" and is followed by a "|---" separator
        if line.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s\-:|]+\|$", lines[i + 1].rstrip()
        ):
            table_lines = []
            while i < len(lines) and lines[i].rstrip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            _render_table(doc, table_lines)
            continue

        # Bullet list
        if line.startswith("- "):
            while i < len(lines) and lines[i].rstrip().startswith("- "):
                item = lines[i].rstrip()[2:]
                # Continuation lines (indented) join onto the same bullet
                i += 1
                while i < len(lines) and lines[i].startswith("  ") and not lines[
                    i
                ].rstrip().startswith("- "):
                    item += " " + lines[i].strip()
                    i += 1
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, item)
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", line):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].rstrip()):
                item = re.sub(r"^\d+\.\s", "", lines[i].rstrip())
                i += 1
                while i < len(lines) and lines[i].startswith("   "):
                    item += " " + lines[i].strip()
                    i += 1
                p = doc.add_paragraph(style="List Number")
                _add_runs(p, item)
            continue

        # Plain paragraph — gather continuation lines until blank
        para_lines = [line]
        i += 1
        while (
            i < len(lines)
            and lines[i].strip()
            and not lines[i].startswith("#")
            and not lines[i].startswith("- ")
            and not lines[i].startswith("|")
            and not lines[i].startswith(">")
            and not re.match(r"^\d+\.\s", lines[i].rstrip())
            and lines[i].strip() != "---"
        ):
            para_lines.append(lines[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        _add_runs(p, " ".join(para_lines))

    doc.save(docx_path)
    print(f"  wrote {docx_path.relative_to(HERE.parent.parent)}")


def _render_table(doc, lines: list[str]) -> None:
    """Render a markdown table into a python-docx Table."""
    # Split each row on "|" and strip; first row is header, second is separator.
    rows = []
    for raw in lines:
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        rows.append(cells)
    header, _separator, *body = rows
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light Grid Accent 1"
    for c, h in enumerate(header):
        cell = table.cell(0, c)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
    for r, row in enumerate(body, start=1):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            _add_runs(cell.paragraphs[0], val)


# --- Email thread ------------------------------------------------------------


def _save_eml(msg: EmailMessage, name: str) -> None:
    path = THREAD_OUT / name
    with path.open("wb") as f:
        f.write(bytes(msg))
    print(f"  wrote {path.relative_to(HERE.parent.parent)}")


def make_email_thread() -> None:
    print("Generating email thread (FOI-2026-058 TrustLab Inc evaluation)...")

    msg_id_1 = "<foi-058-msg1@example.org>"
    msg_id_2 = "<foi-058-msg2@ccsa-bc.example.org>"
    msg_id_3 = "<foi-058-msg3@ccsa-bc.example.org>"
    msg_id_4 = "<foi-058-msg4@ccsa-bc.example.org>"

    # ---- Msg 1: Initial FOI request ----
    m1 = EmailMessage()
    m1["Message-ID"] = msg_id_1
    m1["From"] = f"{REQUESTER[0]} <{REQUESTER[1]}>"
    m1["To"] = f"CCSA FOI Office <{FOI_OFFICER[1]}>"
    m1["Subject"] = "FOI request: TrustLab Inc fractional CIO proposal records"
    m1["Date"] = formatdate(
        datetime(2026, 4, 6, 14, 22, tzinfo=timezone.utc).timestamp(), localtime=False
    )
    m1.set_content(
        f"""\
To the FOI office,

Under the Freedom of Information and Protection of Privacy Act I am
requesting copies of all records held by Coastal Crown Services
Authority concerning its evaluation of the proposal submitted by
TrustLab Inc in response to RFP CCSA-2025-FCIO-014 (Fractional Chief
Information Officer advisory services).

Specifically I am requesting:

  - The original proposal letter and CV submitted by TrustLab Inc.
  - Internal CCSA evaluation memos, scoring records, and committee
    deliberations.
  - Reference-check notes recorded by CCSA staff.
  - Any correspondence between CCSA staff and TrustLab Inc related to
    the procurement.

Personal information about CCSA staff and reference contacts should be
redacted as appropriate under section 22 of FIPPA. I understand that
some committee deliberations may be subject to section 13 (advice and
recommendations) and accept that some content may be withheld on that
basis with a brief explanation.

Please confirm receipt and let me know if a fee estimate applies.

Thank you,
{REQUESTER[0]}
Phone: 555-0145
"""
    )
    _save_eml(m1, "01-foi-request.eml")

    # ---- Msg 2: Acknowledgment ----
    m2 = EmailMessage()
    m2["Message-ID"] = msg_id_2
    m2["From"] = f"{FOI_OFFICER[0]} <{FOI_OFFICER[1]}>"
    m2["To"] = f"{REQUESTER[0]} <{REQUESTER[1]}>"
    m2["Subject"] = "Re: FOI request: TrustLab Inc fractional CIO proposal records"
    m2["Date"] = formatdate(
        datetime(2026, 4, 6, 18, 41, tzinfo=timezone.utc).timestamp(), localtime=False
    )
    m2["In-Reply-To"] = msg_id_1
    m2["References"] = msg_id_1
    m2.set_content(
        f"""\
Hello {REQUESTER[0]},

Thank you for your request, which has been registered as FOI-2026-058.
The 30-business-day clock starts today, April 6, 2026. We anticipate
releasing responsive records by May 20, 2026.

No fees apply at this time. If we determine that the volume of records
or the time required for line-by-line review exceeds the fee threshold
we will contact you with an estimate before processing further.

If you have any questions about the request or its scope, please reply
to this thread.

{FOI_OFFICER[0]}
FOI Coordinator
Coastal Crown Services Authority
"""
    )
    _save_eml(m2, "02-acknowledgment.eml")

    # ---- Msg 3: Internal forward to records custodian ----
    m3 = EmailMessage()
    m3["Message-ID"] = msg_id_3
    m3["From"] = f"{FOI_OFFICER[0]} <{FOI_OFFICER[1]}>"
    m3["To"] = f"Procurement Records <{PROCUREMENT_RECORDS[1]}>"
    m3["Subject"] = "Fwd: FOI request: TrustLab Inc fractional CIO proposal records"
    m3["Date"] = formatdate(
        datetime(2026, 4, 7, 15, 10, tzinfo=timezone.utc).timestamp(), localtime=False
    )
    m3["In-Reply-To"] = msg_id_1
    m3["References"] = msg_id_1
    m3.set_content(
        f"""\
Hi {PROCUREMENT_RECORDS[0]},

Forwarding a new FOI request (FOI-2026-058) covering records related to
the TrustLab Inc proposal under RFP CCSA-2025-FCIO-014. The original
request from {REQUESTER[0]} is quoted below.

Could you pull the following from the procurement file and send them my
way by April 16?

  - Proposal letter and attached CV submitted by TrustLab Inc.
  - Selection Committee evaluation memo (Riley Chen's 17 October memo).
  - Reference-check notes (Riley Chen's 22 October notes).
  - Any TrustLab-related correspondence after 1 October 2025.

Personal information about committee members and reference contacts
will be redacted on this end before release. Send the records as
attachments to your reply.

Thanks,
{FOI_OFFICER[0]}

-------- Forwarded message --------
From: {REQUESTER[0]} <{REQUESTER[1]}>
Date: Mon, Apr 6, 2026 at 2:22 PM UTC
Subject: FOI request: TrustLab Inc fractional CIO proposal records

Under the Freedom of Information and Protection of Privacy Act I am
requesting copies of all records held by Coastal Crown Services
Authority concerning its evaluation of the proposal submitted by
TrustLab Inc in response to RFP CCSA-2025-FCIO-014...
"""
    )
    _save_eml(m3, "03-internal-forward.eml")

    # ---- Msg 4: Internal response with attachments ----
    proposal_path = OUT / "trustlab-proposal-letter.docx"
    memo_path = OUT / "ccsa-evaluation-memo.docx"
    for p in (proposal_path, memo_path):
        if not p.exists():
            raise RuntimeError(
                f"Expected {p.name} to exist by the time the email thread is "
                "built. Generate the DOCX files first."
            )

    m4 = EmailMessage()
    m4["Message-ID"] = msg_id_4
    m4["From"] = f"{PROCUREMENT_RECORDS[0]} <{PROCUREMENT_RECORDS[1]}>"
    m4["To"] = f"{FOI_OFFICER[0]} <{FOI_OFFICER[1]}>"
    m4["Subject"] = "Re: Fwd: FOI request: TrustLab Inc fractional CIO proposal records"
    m4["Date"] = formatdate(
        datetime(2026, 4, 10, 21, 55, tzinfo=timezone.utc).timestamp(), localtime=False
    )
    m4["In-Reply-To"] = msg_id_3
    m4["References"] = f"{msg_id_1} {msg_id_3}"
    m4.set_content(
        f"""\
Hi {FOI_OFFICER[0]},

Attached are the two main records from the TrustLab procurement file:

  - trustlab-proposal-letter.docx — the original 22 September 2025
    proposal letter, including the principal's CV (separate document).
  - ccsa-evaluation-memo.docx — Riley Chen's 17 October internal
    evaluation memo.

The reference-check notes (22 October) are larger and contain a lot of
third-party personal information. I'll send them separately so they
don't trigger our mail relay's attachment-size warning. The CV will
come along with that batch.

Let me know if anything else needs to be pulled from the file.

{PROCUREMENT_RECORDS[0]}
Procurement Records Custodian
Coastal Crown Services Authority
"""
    )
    for path in (proposal_path, memo_path):
        with path.open("rb") as f:
            m4.add_attachment(
                f.read(),
                maintype="application",
                subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=path.name,
            )
    _save_eml(m4, "04-internal-response.eml")


# --- Main --------------------------------------------------------------------


def main() -> int:
    try:
        import docx  # noqa: F401
    except ImportError:
        print(
            "python-docx is required. Install with `pip install python-docx` "
            "or run this script inside the backend container.",
            file=sys.stderr,
        )
        return 1

    print(f"Generating TrustLab/CCSA demo fixtures into {OUT.relative_to(Path.cwd())}/")
    print()

    # DOCX first so the email thread can attach them.
    md_to_docx(SOURCE / "cv-matt-crossley.md", OUT / "trustlab-cv.docx")
    md_to_docx(SOURCE / "proposal-letter.md", OUT / "trustlab-proposal-letter.docx")
    md_to_docx(SOURCE / "evaluation-memo.md", OUT / "ccsa-evaluation-memo.docx")
    md_to_docx(SOURCE / "referee-notes.md", OUT / "ccsa-referee-notes.docx")
    print()

    make_email_thread()
    print()

    print("Done. Outputs:")
    for f in sorted(OUT.rglob("*")):
        if f.is_file():
            print(f"  {f.relative_to(Path.cwd())}  ({f.stat().st_size:>7} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
