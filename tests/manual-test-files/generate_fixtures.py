#!/usr/bin/env python3
"""Generate synthetic test fixtures for manual BlackBar smoke testing.

Run from repo root:
    backend/.venv-test/bin/python3 tests/manual-test-files/generate_fixtures.py

Or from inside the backend container (deps already installed):
    docker compose exec backend python tests/manual-test-files/generate_fixtures.py

Outputs (under tests/manual-test-files/generated/):
  email-thread/
    01-initial-request.eml          Public requester -> FOI office
    02-acknowledgment.eml           FOI office -> requester
    03-internal-forward.eml         FOI officer -> IT dept (forwards #1)
    04-internal-response.eml        IT dept -> FOI officer (with XLSX attachment)
  procurement-policy.docx           Word doc, mixed plain text + a small table
  contractor-list.xlsx              Excel sheet with synthetic vendor data + PII columns
  redaction-test.pdf                Multi-page PDF with names/emails/phones for redaction UI

All content is fabricated. Names, emails, phone numbers, addresses, dollar amounts
are synthetic. The "555-01XX" phone range and "V0V 0V0" Canadian postal code are
explicitly reserved for fictional use.

The email thread uses proper Message-ID / In-Reply-To / References headers so
BlackBar's email-thread consolidation (src/utils/email_threads.py) sees them
as one conversation.
"""
from __future__ import annotations

import os
from datetime import datetime, timezone
from email.message import EmailMessage
from email.utils import formatdate, make_msgid
from pathlib import Path

# --- Output dir --------------------------------------------------------------

ROOT = Path(__file__).parent / "generated"
THREAD_DIR = ROOT / "email-thread"
ROOT.mkdir(exist_ok=True)
THREAD_DIR.mkdir(exist_ok=True)


# --- Synthetic identities ----------------------------------------------------

REQUESTER = ("Jordan Park", "[email protected]")
FOI_OFFICER = ("Riley Chen", "[email protected]")
IT_RECORDS = ("Casey Morgan", "[email protected]")

# 555-01XX is the well-known fictional phone range.
# V0V 0V0 is Canada Post's reserved fictional postal code.

# --- Email thread ------------------------------------------------------------


def _save_eml(msg: EmailMessage, name: str) -> None:
    path = THREAD_DIR / name
    with path.open("wb") as f:
        f.write(bytes(msg))
    print(f"  wrote {path.relative_to(ROOT.parent)}")


def make_email_thread() -> None:
    print("Generating email thread (FOI-2026-042 procurement records)...")

    # Stable Message-IDs so the thread is reproducible across runs.
    msg_id_1 = "<foi-042-msg1@example.org>"
    msg_id_2 = "<foi-042-msg2@blackbar.example>"
    msg_id_3 = "<foi-042-msg3@blackbar.example>"
    msg_id_4 = "<foi-042-msg4@blackbar.example>"

    # ---- Message 1: Initial request (public -> FOI office) ----
    m1 = EmailMessage()
    m1["Message-ID"] = msg_id_1
    m1["From"] = f"{REQUESTER[0]} <{REQUESTER[1]}>"
    m1["To"] = f"FOI Office <{FOI_OFFICER[1]}>"
    m1["Subject"] = "FOI request: Procurement records for IT contracts 2024-2026"
    m1["Date"] = formatdate(
        datetime(2026, 3, 15, 9, 14, tzinfo=timezone.utc).timestamp(), localtime=False
    )
    m1.set_content(
        f"""\
To the FOI office,

Under the Freedom of Information and Protection of Privacy Act I am
requesting copies of all procurement records for information-technology
contracts awarded between January 1, 2024 and March 1, 2026, including:

  - Contract identifiers, vendor names, and total values
  - Award dates and evaluation summaries
  - Signed contracts (with personal information about vendor staff
    redacted as appropriate)

Please confirm receipt and provide a fee estimate if applicable.

Thank you,
{REQUESTER[0]}
Phone: 555-0142
Address: 100 Test Avenue, Anywhere BC V0V 0V0
"""
    )
    _save_eml(m1, "01-initial-request.eml")

    # ---- Message 2: Acknowledgment (FOI office -> requester) ----
    m2 = EmailMessage()
    m2["Message-ID"] = msg_id_2
    m2["From"] = f"{FOI_OFFICER[0]} <{FOI_OFFICER[1]}>"
    m2["To"] = f"{REQUESTER[0]} <{REQUESTER[1]}>"
    m2["Subject"] = "Re: FOI request: Procurement records for IT contracts 2024-2026"
    m2["Date"] = formatdate(
        datetime(2026, 3, 15, 14, 2, tzinfo=timezone.utc).timestamp(), localtime=False
    )
    m2["In-Reply-To"] = msg_id_1
    m2["References"] = msg_id_1
    m2.set_content(
        f"""\
Hello {REQUESTER[0]},

Thank you for your request, which has been registered as FOI-2026-042. The
30-business-day clock starts today, March 15, 2026. We anticipate releasing
responsive records by April 28, 2026.

No fees apply at this time. If we determine that the volume of records
exceeds the fee threshold we will contact you with an estimate before
processing further.

If you have any questions about the request or its scope, please reply to
this thread.

{FOI_OFFICER[0]}
FOI Coordinator
"""
    )
    _save_eml(m2, "02-acknowledgment.eml")

    # ---- Message 3: Internal forward (FOI officer -> IT records) ----
    m3 = EmailMessage()
    m3["Message-ID"] = msg_id_3
    m3["From"] = f"{FOI_OFFICER[0]} <{FOI_OFFICER[1]}>"
    m3["To"] = f"IT Records <{IT_RECORDS[1]}>"
    m3["Subject"] = "Fwd: FOI request: Procurement records for IT contracts 2024-2026"
    m3["Date"] = formatdate(
        datetime(2026, 3, 16, 8, 30, tzinfo=timezone.utc).timestamp(), localtime=False
    )
    m3["In-Reply-To"] = msg_id_1
    m3["References"] = f"{msg_id_1}"
    m3.set_content(
        f"""\
Hi {IT_RECORDS[0]},

Forwarding a new FOI request (FOI-2026-042) for IT procurement records
2024-01-01 through 2026-03-01. The original request from {REQUESTER[0]}
is quoted below.

Could you pull the contract register and any signed contract PDFs for that
period and send them my way by March 25? If signed PDFs include vendor
contact details I'll handle the personal-information redactions on this
end.

Thanks,
{FOI_OFFICER[0]}

-------- Forwarded message --------
From: {REQUESTER[0]} <{REQUESTER[1]}>
Date: Sun, Mar 15, 2026 at 9:14 AM UTC
Subject: FOI request: Procurement records for IT contracts 2024-2026

Under the Freedom of Information and Protection of Privacy Act I am
requesting copies of all procurement records for information-technology
contracts awarded between January 1, 2024 and March 1, 2026...
"""
    )
    _save_eml(m3, "03-internal-forward.eml")

    # ---- Message 4: Internal response with attachment (IT -> FOI officer) ----
    # We attach the XLSX we generate below, so call make_contractor_list_xlsx() first.
    xlsx_path = ROOT / "contractor-list.xlsx"
    if not xlsx_path.exists():
        make_contractor_list_xlsx()
    with xlsx_path.open("rb") as f:
        xlsx_bytes = f.read()

    m4 = EmailMessage()
    m4["Message-ID"] = msg_id_4
    m4["From"] = f"{IT_RECORDS[0]} <{IT_RECORDS[1]}>"
    m4["To"] = f"{FOI_OFFICER[0]} <{FOI_OFFICER[1]}>"
    m4["Subject"] = "Re: Fwd: FOI request: Procurement records for IT contracts 2024-2026"
    m4["Date"] = formatdate(
        datetime(2026, 3, 19, 16, 45, tzinfo=timezone.utc).timestamp(), localtime=False
    )
    m4["In-Reply-To"] = msg_id_3
    m4["References"] = f"{msg_id_1} {msg_id_3}"
    m4.set_content(
        f"""\
Hi {FOI_OFFICER[0]},

Attached is the contractor list for the requested period (Jan 2024 - Mar
2026). 9 active vendor contracts; the Notes column contains some personal
contact information you'll want to review for redaction before release.

Signed contract PDFs are larger files - I'll drop them in the FOI shared
folder under FOI-2026-042/ and notify you when they're ready (probably
March 22).

Let me know if anything needs clarification.

{IT_RECORDS[0]}
IT Records Custodian
"""
    )
    m4.add_attachment(
        xlsx_bytes,
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="contractor-list.xlsx",
    )
    _save_eml(m4, "04-internal-response.eml")


# --- DOCX -------------------------------------------------------------------


def make_procurement_policy_docx() -> None:
    print("Generating procurement-policy.docx...")
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("IT Procurement Policy", level=0)
    p = doc.add_paragraph("Version 2.1  -  Effective March 1, 2026")
    p.runs[0].italic = True

    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph(
        "This policy governs the procurement of information-technology goods "
        "and services for the BlackBar organization. It applies to all "
        "expenditures funded from the IT budget, regardless of source."
    )

    doc.add_heading("2. Authority and Roles", level=1)
    doc.add_paragraph(
        f"The IT Director ({IT_RECORDS[0]}, {IT_RECORDS[1]}) is responsible "
        "for procurement decisions up to the single-vendor limit of "
        "$75,000. The Procurement Committee reviews requests above that "
        "threshold. The FOI Coordinator "
        f"({FOI_OFFICER[0]}, {FOI_OFFICER[1]}) retains all award documents "
        "for the statutory five-year period."
    )

    doc.add_heading("3. Award Thresholds", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    rows = [
        ("Contract value", "Required process"),
        ("Up to $5,000", "Direct award (single quote)"),
        ("$5,001 - $74,999", "Three written quotes; IT Director approval"),
        ("$75,000 and over", "Competitive RFP; Procurement Committee approval"),
    ]
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = val
            if r == 0:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(11)

    doc.add_heading("4. Documentation Requirements", level=1)
    doc.add_paragraph(
        "Each award must be supported by: (a) the original request, (b) the "
        "evaluation matrix, (c) the signed contract, and (d) the procurement "
        "officer's recommendation memo. These records are subject to FOI "
        "release under the standard redaction schedule (personal identifiers "
        "redacted; contract values and vendor names released)."
    )

    doc.add_heading("5. Conflict of Interest", level=1)
    doc.add_paragraph(
        "Procurement officers must declare any personal or financial "
        "relationship with a bidding vendor before any evaluation activity. "
        "Declarations are retained by the FOI Coordinator and are not "
        "themselves subject to release."
    )

    doc.add_heading("6. Contact", level=1)
    doc.add_paragraph(
        f"Questions about this policy should be directed to {IT_RECORDS[0]} "
        f"at {IT_RECORDS[1]} or via the procurement intake line at "
        "555-0144."
    )

    out = ROOT / "procurement-policy.docx"
    doc.save(out)
    print(f"  wrote {out.relative_to(ROOT.parent)}")


# --- XLSX --------------------------------------------------------------------


def make_contractor_list_xlsx() -> None:
    print("Generating contractor-list.xlsx...")
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "IT Contracts 2024-2026"

    header = [
        "Vendor ID",
        "Vendor Name",
        "Primary Contact",
        "Contact Email",
        "Contact Phone",
        "Contract Value",
        "Award Date",
        "Notes",
    ]
    for col, h in enumerate(header, start=1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")

    rows = [
        ("V-1001", "Acme Cloud Systems", "Dana Whitfield",
         "[email protected]", "555-0101", 42500.00, "2024-04-12",
         "Annual SaaS renewal. Primary contact reachable Mon-Fri 8-4 PT."),
        ("V-1002", "Northwind Networks Ltd", "Avery Sutton",
         "[email protected]", "555-0102", 118200.00, "2024-06-30",
         "Network refresh contract. Sutton noted on-call after hours."),
        ("V-1003", "Helix Cybersecurity", "Morgan Avery",
         "[email protected]", "555-0103", 86000.00, "2024-09-17",
         "Pentest + ongoing managed-SOC. Auto-renewing."),
        ("V-1004", "Brightspark Web Co.", "Quinn Halloran",
         "[email protected]", "555-0104", 18750.00, "2024-11-04",
         "Web redesign one-time. Halloran is contractor's sole staff."),
        ("V-1005", "Mira Hardware", "Pat Iversen",
         "[email protected]", "555-0105", 9800.00, "2025-01-08",
         "Misc peripherals. Iversen home phone for after-hours: 555-0179."),
        ("V-1006", "QuantumLeap Analytics", "Sam Rivera",
         "[email protected]", "555-0106", 215000.00, "2025-03-22",
         "Two-year BI engagement. Rivera also serves as account exec for V-1010."),
        ("V-1007", "Pinewood IT Staffing", "Robin Chase",
         "[email protected]", "555-0107", 64200.00, "2025-06-11",
         "Augmented staffing for migration project. Chase placed 3 contractors."),
        ("V-1008", "AcorNet Hosting", "Cory Beckman",
         "[email protected]", "555-0108", 31000.00, "2025-08-29",
         "Hosting for public records portal."),
        ("V-1009", "Lighthouse Print & Scan", "Drew Kalvinski",
         "[email protected]", "555-0109", 7400.00, "2025-12-15",
         "Records-scan project. Kalvinski's mobile: 555-0185."),
    ]
    for r, row_data in enumerate(rows, start=2):
        for c, val in enumerate(row_data, start=1):
            cell = ws.cell(row=r, column=c, value=val)
            if c == 6:  # Contract Value column
                cell.number_format = '"$"#,##0.00'

    # Reasonable column widths
    widths = [10, 26, 22, 32, 14, 16, 12, 60]
    for col_idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + col_idx)].width = width

    out = ROOT / "contractor-list.xlsx"
    wb.save(out)
    print(f"  wrote {out.relative_to(ROOT.parent)}")


# --- PDF ---------------------------------------------------------------------


def make_redaction_test_pdf() -> None:
    print("Generating redaction-test.pdf...")
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        PageBreak,
    )

    out = ROOT / "redaction-test.pdf"
    doc = SimpleDocTemplate(
        str(out),
        pagesize=LETTER,
        title="FOI-2026-042 Briefing Note",
        author="Synthetic test fixture",
    )
    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    body = styles["BodyText"]

    flow = []
    flow.append(Paragraph("Briefing Note: FOI-2026-042", h1))
    flow.append(
        Paragraph(
            "Prepared by Riley Chen, FOI Coordinator. Date: 2026-03-25. "
            "Status: pre-release draft.",
            body,
        )
    )
    flow.append(Spacer(1, 12))

    flow.append(Paragraph("1. Request Summary", h2))
    flow.append(
        Paragraph(
            "On March 15, 2026, this office received an FOI request from "
            "Jordan Park (100 Test Avenue, Anywhere BC V0V 0V0; phone "
            "555-0142; email [email protected]) seeking procurement "
            "records for IT contracts awarded between 2024-01-01 and "
            "2026-03-01. The request was acknowledged the same day and "
            "logged as FOI-2026-042.",
            body,
        )
    )

    flow.append(Paragraph("2. Records Located", h2))
    flow.append(
        Paragraph(
            "IT Records (custodian: Casey Morgan, [email protected], "
            "555-0144) identified 9 contracts matching the search criteria. "
            "The contractor list and signed PDFs have been transferred to "
            "the FOI working folder.",
            body,
        )
    )

    flow.append(Paragraph("3. Personal Information in Records", h2))
    flow.append(
        Paragraph(
            "The following personal identifiers appear in the responsive "
            "records and are recommended for redaction under FIPPA s.22 "
            "(unreasonable invasion of personal privacy):",
            body,
        )
    )
    flow.append(
        Paragraph(
            "&nbsp;&nbsp;- Vendor contact mobile numbers (e.g., Pat Iversen "
            "after-hours 555-0179; Drew Kalvinski mobile 555-0185).",
            body,
        )
    )
    flow.append(
        Paragraph(
            "&nbsp;&nbsp;- Personal email addresses where contractors used "
            "non-corporate domains (none in this batch).",
            body,
        )
    )
    flow.append(
        Paragraph(
            "&nbsp;&nbsp;- A handwritten note on the Helix Cybersecurity "
            "contract showing the home address of Morgan Avery "
            "(42 Practice Lane, Anywhere BC V0V 0V1).",
            body,
        )
    )

    flow.append(Paragraph("4. Recommendations", h2))
    flow.append(
        Paragraph(
            "Release contract identifiers, vendor names, dollar values, and "
            "award dates without redaction. Redact the personal identifiers "
            "listed above. Provide a brief response letter referencing "
            "FIPPA s.22.",
            body,
        )
    )

    flow.append(Paragraph("5. Sign-off", h2))
    flow.append(
        Paragraph(
            "Reviewer: Riley Chen, FOI Coordinator. "
            "Phone: 555-0143. Email: [email protected]. "
            "Recommendation: release with redactions as noted above.",
            body,
        )
    )

    flow.append(PageBreak())
    flow.append(Paragraph("Appendix A: Vendor Notes (Excerpt)", h1))
    flow.append(Spacer(1, 8))
    flow.append(
        Paragraph(
            "Excerpted from the contractor-list.xlsx Notes column. Reproduced "
            "here only for redaction-marking convenience; the spreadsheet is "
            "the authoritative source.",
            body,
        )
    )
    flow.append(Spacer(1, 8))
    vendor_notes = [
        "V-1001 Acme Cloud Systems / Dana Whitfield ([email protected], "
        "555-0101): Annual SaaS renewal. Primary contact reachable Mon-Fri "
        "8-4 PT.",
        "V-1005 Mira Hardware / Pat Iversen ([email protected], "
        "555-0105): Misc peripherals. Iversen home phone for after-hours: "
        "555-0179.",
        "V-1009 Lighthouse Print & Scan / Drew Kalvinski "
        "([email protected], 555-0109): Records-scan project. "
        "Kalvinski's mobile: 555-0185.",
    ]
    for n in vendor_notes:
        flow.append(Paragraph(n, body))
        flow.append(Spacer(1, 6))

    doc.build(flow)
    print(f"  wrote {out.relative_to(ROOT.parent)}")


# --- Main --------------------------------------------------------------------


def main() -> None:
    print(f"Generating fixtures into {ROOT.relative_to(Path.cwd())}/")
    make_contractor_list_xlsx()
    make_procurement_policy_docx()
    make_redaction_test_pdf()
    make_email_thread()
    print()
    print("Done. Outputs:")
    for f in sorted(ROOT.rglob("*")):
        if f.is_file():
            print(f"  {f.relative_to(Path.cwd())}  ({f.stat().st_size:>7} bytes)")


if __name__ == "__main__":
    main()
